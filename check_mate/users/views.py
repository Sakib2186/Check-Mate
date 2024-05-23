from django.shortcuts import render,HttpResponse,redirect,reverse
import logging
from datetime import datetime
import traceback
from django.contrib.auth.models import auth,User
from django.contrib import messages
from system_administrator.models import *
from system_administrator.system_error_handling import ErrorHandling
from .render_data import *
from .models import *
import pyotp
from django.utils.safestring import mark_safe
from django.contrib.auth.decorators import login_required
from check_mate import settings
from docx import Document
from docx.shared import Pt,Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from io import BytesIO
from django.http import JsonResponse
import xlwt
import base64
import os,time
from django.core.files.base import ContentFile

logger=logging.getLogger(__name__)

# Create your views here.

def check_mate(request):
    if request.user.is_authenticated:
        return redirect('users:dashboard')
    else:
        return redirect('users:login')

    
def login(request):

    try:

        if request.method == "POST":

            if request.POST.get('login_button_pressed'):
                #when user presses login in button fetching the data he gave

                user_username = request.POST.get('user_username')
                user_password = request.POST.get('user_password')
                
                #getting user credentials from django users model
                #it would be valid for both admin/superuser and normal users of system
                user = auth.authenticate(username=user_username,password=user_password)
                #if user not found then logged him in
                if user is not None:
                    
                    try:
                        school_user = School_Users.objects.get(user_id = user_username)
                        #checking to see if school user is verified or not
                        if school_user.user_otp_verified:
                            auth.login(request,user)
                            return redirect('users:dashboard')
                        else:
                            #not verified so sending them the link for verification
                            verification_link = reverse('users:registration_email_verification', args=[school_user.user_id])
                            verification_url = request.build_absolute_uri(verification_link)
                            verification_message = f"Your account is not verified yet! Verification link: <a href='{verification_url}'>{verification_url}</a>"
                            messages.error(request, mark_safe(verification_message))
                            return redirect('users:login') 
                    except:
                        #admin is logging in
                        auth.login(request,user)
                        return redirect('users:dashboard')
                else:
                    #user might not be registered or gives the wrong credentials
                    messages.info(request,"Credentials given are wrong")
                    return redirect('users:login') 

        context = {
            'page_title':"Check Mate"
        }
        
        return render(request,"login_page.html",context)
    
    except Exception as e:
        #saving error information in database if error occured
        logger.error("An error occurred for during logging in at {datetime}".format(datetime=datetime.now()), exc_info=True)
        ErrorHandling.save_system_errors('Login Error',error_name=e,error_traceback=traceback.format_exc())
        return HttpResponse("Bad Request")
    
def registration(request):

    try:

        if request.method == "POST":

            if request.POST.get('registration_button'):
                
                #getting user credentials
                user_name = request.POST.get('user_name')
                password = request.POST.get('user_password')
                confirm_password = request.POST.get('confirm_user_password')

                #getting maipulated data from render_data
                registration = Login.register_user(request,user_name,password,confirm_password)
                #checking to see if user is successfully registered
                if registration[0]:
                    return redirect('users:registration_email_verification',registration[1].user_id)
                else:
                    messages.error(request,registration[1])
                    return redirect('users:registration')

        context = {
            'page_title':'Check Mate'
        }

        return render(request,"registration.html",context)

    except Exception as e:
        #saving error information in database if error occured
        logger.error("An error occurred for during logging in at {datetime}".format(datetime=datetime.now()), exc_info=True)
        ErrorHandling.save_system_errors('Registration Error',error_name=e,error_traceback=traceback.format_exc())
        return HttpResponse("Bad Request")

def registration_email_verification(request,user_id):

    try:

        school_user = School_Users.objects.get(user_id = user_id)
        #if user is not verified then allowing to enter page
        if not school_user.user_otp_verified:

            if request.method == "POST":

                if request.POST.get('submit_otp'):
                    
                    #getting user otp
                    otp = str(request.POST.get('otp'))
                    
                    #checking if otp is empty
                    if otp == "":
                        messages.error(request,'Provide OTP!')
                        return redirect('users:registration_email_verification',user_id)
                    try:
                        #putting them session dictionary
                        otp_secret_key = request.session['otp_secret_key']
                        otp_valid_date = request.session['otp_valid_date']
                        print("here")
                        #checking to see if otp is valid and matches with users input one
                        if otp_secret_key and otp_valid_date is not None:
                            valid_until = datetime.fromisoformat(otp_valid_date)

                            if valid_until > datetime.now():
                                totp = pyotp.TOTP(otp_secret_key,interval = 60)
                                print("here2")
                                if totp.verify(otp):
                                    user = User.objects.get(username = user_id)
                                    school_user.user_otp_verified = True
                                    school_user.save()
                                    auth.login(request,user)
                                    #removing the otp keys from the session
                                    request.session.pop('otp_secret_key', None)
                                    request.session.pop('otp_valid_date', None)
                                    return redirect('users:dashboard')
                                else:
                                    messages.error(request,"invalid OTP")
                                    return redirect('users:registration_email_verification',user_id)
                            else:
                                messages.error(request,"OTP expired")
                                redirect('users:registration_email_verification',user_id)

                        else:
                            messages.error(request,"Something went wrong..")
                            redirect('users:registration_email_verification',user_id)
                    except:
                        messages.error(request,"invalid OTP")
                        return redirect('users:registration_email_verification',user_id)

                if request.POST.get('resend_otp'):
                    
                    request.session.pop('otp_secret_key', None)
                    request.session.pop('otp_valid_date', None)
                    Login.registration_email(request,school_user)
                    redirect('users:registration_email_verification',user_id)
                    
            context = {
                'page_title':'Check Mate'
            }
            return render(request,"registration_verification.html",context)
        
        else:
            return HttpResponse("Already verified")

    except Exception as e:
        #saving error information in database if error occured
        logger.error("An error occurred for during logging in at {datetime}".format(datetime=datetime.now()), exc_info=True)
        ErrorHandling.save_system_errors('Registration Error',error_name=e,error_traceback=traceback.format_exc())
        return HttpResponse("Bad Request")

@login_required
def dashboard(request):

    try:
        #loading the data to pass them in dictionary, context
        type_of_logged_in_user = Login.user_type_logged_in(request)
        logged_in_user = Login.logged_in_user(request)
        current_semester = Session.objects.get(current=True)
        bar_chart_stats = Load_Courses.get_bar_chart_stats(logged_in_user)
        print(bar_chart_stats[0])
        print(bar_chart_stats[1])
        print(bar_chart_stats[2])
        print(bar_chart_stats[3])
        if logged_in_user == None:
            user = request.user.username
        else:
            user = logged_in_user.user_id
        context = {
            'page_title':'Check Mate',
            'user_type':type_of_logged_in_user,
            'media_url':settings.MEDIA_URL,
            'logged_in_user':logged_in_user,
            'year':datetime.now().year,
            'current_semester':current_semester,
            'bar_chart_stats':bar_chart_stats,
        }

        return render(request,"dashboard.html",context)

    except Exception as e:
        #saving error information in database if error occured
        logger.error("An error occurred for during logging in at {datetime}".format(datetime=datetime.now()), exc_info=True)
        ErrorHandling.save_system_errors(user,error_name=e,error_traceback=traceback.format_exc())
        return HttpResponse("Bad Request")
    
@login_required
def logout(request):
    
    try:
        auth.logout(request)
        return redirect('users:login')
    except Exception as e:
        #saving error information in database if error occured
        logger.error("An error occurred for during logging in at {datetime}".format(datetime=datetime.now()), exc_info=True)
        ErrorHandling.save_system_errors('Logout',error_name=e,error_traceback=traceback.format_exc())
        return HttpResponse("Bad Request")

@login_required
def edit_profile(request):

    try:
        #loading the data to pass them in dictionary, context
        type_of_logged_in_user = Login.user_type_logged_in(request)
        logged_in_user = Login.logged_in_user(request)
        current_semester = Session.objects.get(current=True)

        if logged_in_user == None:
            user = request.user.username
        else:
            user = logged_in_user.user_id
        context = {
            'page_title':'Check Mate',
            'user_type':type_of_logged_in_user,
            'media_url':settings.MEDIA_URL,
            'logged_in_user':logged_in_user,
            'year':datetime.now().year,
            'current_semester':current_semester,
        }
        return render(request,"edit_account.html",context)

    except Exception as e:
        #saving error information in database if error occured
        logger.error("An error occurred for during logging in at {datetime}".format(datetime=datetime.now()), exc_info=True)
        ErrorHandling.save_system_errors(user,error_name=e,error_traceback=traceback.format_exc())
        return HttpResponse("Bad Request")
    
@login_required
def courses(request):

    try:
        #loading the data to pass them in dictionary, context
        type_of_logged_in_user = Login.user_type_logged_in(request)
        logged_in_user = Login.logged_in_user(request)
        all_courses = Load_Courses.get_user_courses(logged_in_user)
        current_semester = Session.objects.get(current=True)

        if logged_in_user == None:
            user = request.user.username
        else:
            user = logged_in_user.user_id
        context = {
            'page_title':'Check Mate',
            'user_type':type_of_logged_in_user,
            'media_url':settings.MEDIA_URL,
            'logged_in_user':logged_in_user,
            'year':datetime.now().year,
            'current_semester':current_semester,

            'all_courses':all_courses,
        }

        return render(request,"courses.html",context)


    except Exception as e:
        #saving error information in database if error occured
        logger.error("An error occurred for during logging in at {datetime}".format(datetime=datetime.now()), exc_info=True)
        ErrorHandling.save_system_errors(user,error_name=e,error_traceback=traceback.format_exc())
        return HttpResponse("Bad Request")
    
@login_required
def all_courses(request):

    try:
        #loading the data to pass them in dictionary, context
        type_of_logged_in_user = Login.user_type_logged_in(request)
        logged_in_user = Login.logged_in_user(request)
        all_courses = Load_Courses.get_user_courses(logged_in_user)
        current_semester = Session.objects.get(current=True)
        

        if logged_in_user == None:
            user = request.user.username
        else:
            user = logged_in_user.user_id

        if logged_in_user == None:
            #as admin so loading all courses 
            courses_all = Load_Courses.get_all_courses()
            
            if request.POST.get('save_section'):

                section_number = int(request.POST.get('section_number'))
                course = request.POST.get('course_code')

                result = Save.save_section(section_number,course)
                if result[0]:
                    messages.success(request,result[1])
                    return redirect('users:all_courses')
                else:
                    messages.error(request,result[1])
                    return redirect('users:all_courses')
                

            context = {
                'page_title':'Check Mate',
                'user_type':type_of_logged_in_user,
                'media_url':settings.MEDIA_URL,
                'logged_in_user':logged_in_user,
                'year':datetime.now().year,
                'current_semester':current_semester,

                'all_courses':all_courses,
                'courses':courses_all,

            }

            return render(request,"all_courses.html",context)
        else:
            context = {
                    'page_title':'Check Mate',
                    'user_type':type_of_logged_in_user,
                    'media_url':settings.MEDIA_URL,
                    'logged_in_user':logged_in_user,
                    'year':datetime.now().year,
                    'current_semester':current_semester,
            }
            return render(request,"access_denied.html",context)

    except Exception as e:
        #saving error information in database if error occured
        logger.error("An error occurred for during logging in at {datetime}".format(datetime=datetime.now()), exc_info=True)
        ErrorHandling.save_system_errors(user,error_name=e,error_traceback=traceback.format_exc())
        return HttpResponse("Bad Request")

@login_required
def add_course(request):

    try:
        #loading the data to pass them in dictionary, context
        type_of_logged_in_user = Login.user_type_logged_in(request)
        logged_in_user = Login.logged_in_user(request)
        current_semester = Session.objects.get(current=True)

        if logged_in_user == None:
            user = request.user.username
        else:
            user = logged_in_user.user_id

        if logged_in_user == None:

            if request.method == "POST":

                if request.POST.get('save_course'):

                    course_name = request.POST.get('course_name')
                    course_code = request.POST.get('course_code')
                    course_description = request.POST.get('course_description')
                    try:
                        cover_picture = request.FILES['cover_picture']
                    except:
                        cover_picture = None
                    
                    result = Save.save_course(course_code,course_name,course_description,cover_picture,course_existing = None)
                    if result[0]:
                        messages.success(request,'Course Added Successfully!')
                        return redirect('users:edit_course_details',result[1].pk)
                    else:
                        messages.error(request,"Something went wrong. Try again later")
                        return redirect('users:add_course')

                    
            context = {
                'page_title':'Check Mate',
                'user_type':type_of_logged_in_user,
                'media_url':settings.MEDIA_URL,
                'logged_in_user':logged_in_user,
                'year':datetime.now().year,
                'current_semester':current_semester,

            }
            return render(request,"course_edit.html",context)
        else:
            context = {
                    'page_title':'Check Mate',
                    'user_type':type_of_logged_in_user,
                    'media_url':settings.MEDIA_URL,
                    'logged_in_user':logged_in_user,
                    'year':datetime.now().year,
                    'current_semester':current_semester,
            }
            return render(request,"access_denied.html",context)

    except Exception as e:
        #saving error information in database if error occured
        logger.error("An error occurred for during logging in at {datetime}".format(datetime=datetime.now()), exc_info=True)
        ErrorHandling.save_system_errors(user,error_name=e,error_traceback=traceback.format_exc())
        return HttpResponse("Bad Request")
    
@login_required
def edit_course_details(request,course_id):

    try:
        #loading the data to pass them in dictionary, context
        type_of_logged_in_user = Login.user_type_logged_in(request)
        logged_in_user = Login.logged_in_user(request)
        current_semester = Session.objects.get(current=True)

        if logged_in_user == None:
            user = request.user.username
        else:
            user = logged_in_user.user_id

        if logged_in_user == None:
            course = Load_Courses.get_specific_course(course_id)

            if request.method == "POST":

                if request.POST.get('save_course'):

                    course_name = request.POST.get('course_name')
                    course_code = request.POST.get('course_code')
                    course_description = request.POST.get('course_description')
                    cover_picture = request.FILES.get('cover_picture')
                    if cover_picture == "":
                        cover_picture = None
                    result = Save.save_course(course_code,course_name,course_description,cover_picture,course)
                    if result[0]:
                        messages.success(request,'Course Updated Successfully!')
                        return redirect('users:edit_course_details',result[1].pk)
                    else:
                        messages.error(request,"Something went wrong. Try again later")
                        return redirect('users:edit_course_details',result[1].pk)
                    
                if request.POST.get('delete_course_flag'):

                    flag = str(request.POST.get('delete_course_flag'))
                    if flag == str(1):
                        
                        if Delete.delete_course(course_id):
                            messages.success(request,"Course deleted successfully!")
                            return redirect('users:all_courses')
                        else:
                            messages.error(request,"Could not delete. Please try again.")
                            return redirect('users:edit_course_details',course_id)
            context = {
                'page_title':'Check Mate',
                'user_type':type_of_logged_in_user,
                'media_url':settings.MEDIA_URL,
                'logged_in_user':logged_in_user,
                'year':datetime.now().year,
                'current_semester':current_semester,

                'course':course,

            }
            return render(request,"course_edit.html",context)
        else:
            context = {
                    'page_title':'Check Mate',
                    'user_type':type_of_logged_in_user,
                    'media_url':settings.MEDIA_URL,
                    'logged_in_user':logged_in_user,
                    'year':datetime.now().year,
                    'current_semester':current_semester,
            }
            return render(request,"access_denied.html",context)


    except Exception as e:
        #saving error information in database if error occured
        logger.error("An error occurred for during logging in at {datetime}".format(datetime=datetime.now()), exc_info=True)
        ErrorHandling.save_system_errors(user,error_name=e,error_traceback=traceback.format_exc())
        return HttpResponse("Bad Request")

@login_required
def save_semester(request):

    try:
        #loading the data to pass them in dictionary, context
        type_of_logged_in_user = Login.user_type_logged_in(request)
        logged_in_user = Login.logged_in_user(request)

        if logged_in_user == None:
            user = request.user.username
        else:
            user = logged_in_user.user_id

        if request.method == 'POST':
            if request.POST.get('semeter_submit'):
                selected_semester = int(request.POST.get('selectedSemester'))
                checkbox_checked = request.POST.get('checkboxChecked')
                year = int(request.POST.get('year'))

                if selected_semester == 0:
                    messages.error(request,'Could not save!')
                    return redirect('users:dashboard')

                if checkbox_checked == "on":
                    
                    if Save.save_semester(selected_semester,year):
                        messages.success(request,'Semester Saved!')
                        return redirect('users:dashboard')
                    else:
                        messages.error(request,'Could not save!')
                        return redirect('users:dashboard')

        context = {
            'page_title':'Check Mate',
            'user_type':type_of_logged_in_user,
            'media_url':settings.MEDIA_URL,
            'logged_in_user':logged_in_user,
            'year':datetime.now().year,
        }

        return render(request,"dashboard.html",context)

    except Exception as e:
        #saving error information in database if error occured
        logger.error("An error occurred for during logging in at {datetime}".format(datetime=datetime.now()), exc_info=True)
        ErrorHandling.save_system_errors(user,error_name=e,error_traceback=traceback.format_exc())
        return HttpResponse("Bad Request")
        

@login_required
def course_edit(request,course_id):

    try:
        #loading the data to pass them in dictionary, context
        type_of_logged_in_user = Login.user_type_logged_in(request)
        logged_in_user = Login.logged_in_user(request)
        current_semester = Session.objects.get(current=True)

        if logged_in_user == None:
            user = request.user.username
        else:
            user = logged_in_user.user_id

        if logged_in_user == None:

            selected_members = Load_Courses.get_selected_ta_students_instructors(course_id)    

            if request.method == "POST":

                if request.POST.get('save_course_section'):

                    course = Course_Section.objects.get(id = course_id)
                    section_number = course.section_number
                    instructor = request.POST.get('instructor_checbox')
                    ta = request.POST.get('ta_checkbox')
                    students = request.POST.getlist('student_checkbox')

                    result =  Save.save_course_section_details(section_number,instructor,ta,students,course_id,selected_members[3])
                    if result[0]:
                        messages.success(request,result[1])
                        return redirect('users:course_edit',course_id)
                    else:
                        messages.error(request,"Could not save!")
                        return redirect('users:course_edit',course_id)

            instructors = Load_Courses.get_all_instructors(selected_members[0])
            students = Load_Courses.get_all_student(selected_members[2])
            course_details = Load_Courses.get_specific_course_section(course_id)

            context = {
                'page_title':'Check Mate',
                'user_type':type_of_logged_in_user,
                'media_url':settings.MEDIA_URL,
                'logged_in_user':logged_in_user,
                'year':datetime.now().year,
                'current_semester':current_semester,

                'course_details':course_details,
                'all_instructors':instructors,
                'all_students':students,
                'selected_members':selected_members,
                'exist':selected_members[3]
            }

            return render(request,"course_edit2.html",context)
        else:
            context = {
                    'page_title':'Check Mate',
                    'user_type':type_of_logged_in_user,
                    'media_url':settings.MEDIA_URL,
                    'logged_in_user':logged_in_user,
                    'year':datetime.now().year,
                    'current_semester':current_semester,
            }
            return render(request,"access_denied.html",context)

    except Exception as e:
        #saving error information in database if error occured
        logger.error("An error occurred for during logging in at {datetime}".format(datetime=datetime.now()), exc_info=True)
        ErrorHandling.save_system_errors(user,error_name=e,error_traceback=traceback.format_exc())
        return HttpResponse("Bad Request")
@login_required            
def course(request,course_id):

    try:
         #loading the data to pass them in dictionary, context
        type_of_logged_in_user = Login.user_type_logged_in(request)
        logged_in_user = Login.logged_in_user(request)
        current_semester = Session.objects.get(current=True)
        section_exams = Load_Courses.get_section_exams(course_id)
        exam =Load_Courses.get_specific_course_section(course_id)
        logged_in_ta = False
        
        try:
            submitted_by = exam.teaching_assistant.all()
            ta = submitted_by[0].teaching_id
            if ta == logged_in_user:
                logged_in_ta = True
        except:
            pass
        

        if logged_in_user == None:
            user = request.user.username
        else:
            user = logged_in_user.user_id


        context = {
                'page_title':'Check Mate',
                'user_type':type_of_logged_in_user,
                'media_url':settings.MEDIA_URL,
                'logged_in_user':logged_in_user,
                'year':datetime.now().year,
                'current_semester':current_semester,

                'course_id':course_id,
                'section_exams':section_exams,
                'logged_in_ta':logged_in_ta,
            }
        
        return render(request,"exam_homepage.html",context)
    except Exception as e:
        #saving error information in database if error occured
        logger.error("An error occurred for during logging in at {datetime}".format(datetime=datetime.now()), exc_info=True)
        ErrorHandling.save_system_errors(user,error_name=e,error_traceback=traceback.format_exc())
        return HttpResponse("Bad Request") 
@login_required
def take_exam(request,course_id):

    try:
         #loading the data to pass them in dictionary, context
        type_of_logged_in_user = Login.user_type_logged_in(request)
        logged_in_user = Login.logged_in_user(request)
        current_semester = Session.objects.get(current=True)
        exam_modes = Exam_Mode.objects.all()
        exam_type = Exam_Type.objects.all()
    

        if logged_in_user == None:
            user = request.user.username
        else:
            user = logged_in_user.user_id
    
        if 1 in type_of_logged_in_user or 3 in type_of_logged_in_user:

            if request.method == "POST":

                if request.POST.get('save_exam'):

                    exam_title = request.POST.get('exam_title')
                    exam_type = int(request.POST.get('exam_type'))
                    exam_mode = int(request.POST.get('exam_mode'))
                    exam_date = request.POST.get('exam_date')
                    exam_description = request.POST.get('exam_description')
                    exam_set = request.POST.get('exam_set')
        
                    if exam_set == "":
                        exam_set=0

                    if exam_mode == 0:
                        messages.error(request,"Please select Exam Mode!")
                        return redirect('users:take_exam',course_id)
                    if exam_type == 0:
                        messages.error(request,"Please select Exam Type!")
                        return redirect('users:take_exam',course_id)

                    if exam_date == "":
                        messages.error(request,"Please provide Exam date!")
                        return redirect('users:take_exam',course_id)

                    result= Save.save_exams_for_section(course_id,exam_title,exam_type,exam_mode,exam_date,exam_description,exam_set,exam_id=None,ta_available=False)
                    if result[0]:
                        messages.success(request,result[1])
                        return redirect('users:edit_exam',course_id,result[2].pk)
                    else:
                        messages.error(request,"Could not save! Try again!")
                        return redirect('users:take_exam',course_id)

            context = {
                    'page_title':'Check Mate',
                    'user_type':type_of_logged_in_user,
                    'media_url':settings.MEDIA_URL,
                    'logged_in_user':logged_in_user,
                    'year':datetime.now().year,
                    'current_semester':current_semester,

                    'course_id':course_id,
                    'exam_modes':exam_modes,
                    'exam_types':exam_type,

                }
            
            return render(request,"add_exam.html",context)
        else:
            context = {
                    'page_title':'Check Mate',
                    'user_type':type_of_logged_in_user,
                    'media_url':settings.MEDIA_URL,
                    'logged_in_user':logged_in_user,
                    'year':datetime.now().year,
                    'current_semester':current_semester,
            }
            return render(request,"access_denied.html",context)
    except Exception as e:
        #saving error information in database if error occured
        logger.error("An error occurred for during logging in at {datetime}".format(datetime=datetime.now()), exc_info=True)
        ErrorHandling.save_system_errors(user,error_name=e,error_traceback=traceback.format_exc())
        return HttpResponse("Bad Request") 

@login_required
def edit_exam(request,course_id,exam_id):

    try:
        #loading the data to pass them in dictionary, context
        type_of_logged_in_user = Login.user_type_logged_in(request)
        logged_in_user = Login.logged_in_user(request)
        current_semester = Session.objects.get(current=True)
        exam_modes = Exam_Mode.objects.all()
        exam_type = Exam_Type.objects.all()
        section_exam = Load_Courses.get_saved_section_exams(exam_id)
        try:
            announce = Announcements.objects.get(section_exam=section_exam)
            print(announce)
        except:
            announce = None
        try:
            time = section_exam.exam_time.split(" ")
        except:
            time = None


        if logged_in_user == None:
            user = request.user.username
        else:
            user = logged_in_user.user_id
        

        if 1 in type_of_logged_in_user or 3 in type_of_logged_in_user:

            if request.method == "POST":

                if request.POST.get('save_exam'):

                    exam_title = request.POST.get('exam_title')
                    exam_type = int(request.POST.get('exam_type'))
                    exam_mode = int(request.POST.get('exam_mode'))
                    exam_date = request.POST.get('exam_date')
                    exam_description = request.POST.get('exam_description')
                    exam_set = request.POST.get('exam_set')
                    ta_available = request.POST.get('ta_available')

                    if ta_available == "on":
                        ta_available= True
                    else:
                        ta_available=False

                    if exam_set == "" or exam_set=="0":
                        exam_set=0
                    
                    if exam_set=="1":
                        messages.error(request,"Exam set must be 2 or 3 !")
                        return redirect('users:edit_exam',course_id,exam_id)

                    if exam_mode == 0:
                        messages.error(request,"Please select Exam Mode!")
                        return redirect('users:edit_exam',course_id,exam_id)
                    if exam_type == 0:
                        messages.error(request,"Please select Exam Type!")
                        return redirect('users:edit_exam',course_id,exam_id)

                    if exam_date == "":
                        messages.error(request,"Please provide Exam date!")
                        return redirect('users:edit_exam',course_id,exam_id)

                    result= Save.save_exams_for_section(course_id,exam_title,exam_type,exam_mode,exam_date,exam_description,exam_set,exam_id,ta_available)
                    if result[0]:
                        messages.success(request,result[1])
                        return redirect('users:edit_exam',course_id,result[2].pk)
                    else:
                        messages.error(request,"Could not save! Try again!")
                        return redirect('users:edit_exam',course_id,result[2].pk)

                if request.POST.get('save_question'):

                    question_set = request.POST.get('question_set')
                    question = request.POST.get('question')
                    answer_size = request.POST.get('answer_size')
                    marks = request.POST.get('marks')
                    box_height = int(request.POST.get('box_height'))
                    question_image = request.FILES.get('question_image_add')


                    if question_image == "":
                        question_image = None

                    if answer_size == "0":
                        messages.error(request,"Select Answer Field!")
                        return redirect('users:edit_exam',course_id,section_exam.pk)
                    
                    if box_height == 0:
                        messages.error(request,"Box Height cannot be 0!")
                        return redirect('users:edit_exam',course_id,section_exam.pk)

                    result = Save.save_question_for_exam(exam_id,question_set,question,answer_size,marks,box_height,question_image,question_id=None)

                    if result[0]:
                        messages.success(request,result[1])
                        return redirect('users:edit_exam',course_id,section_exam.pk)
                    else:
                        messages.error(request,"Error occured! Try again later.")
                        return redirect('users:edit_exam',course_id,section_exam.pk)
                    
                if request.POST.get('delete_question'):
                        
                        flag = str(request.POST.get('delete_question'))
                        pk = request.POST.get('question_pk')
        
                        if flag == "1":

                            if Delete.delete_question(pk):
                                messages.success(request,"Question Deleted Successfully!")
                                return redirect('users:edit_exam',course_id,section_exam.pk)
                            else:
                                messages.error(request,"Error occured! Try again later.")
                                return redirect('users:edit_exam',course_id,section_exam.pk)
                            
                if request.POST.get('edit_question'):

                    question_pk = request.POST.get('question_pk')
                    question_set = request.POST.get('question_set')
                    question = request.POST.get('question')
                    answer_size = request.POST.get('answer_size2')
                    marks = request.POST.get('marks')
                    box_height = int(request.POST.get('box_height2'))
                    question_image = request.FILES.get('question_image_edit')

                    if question_image == "":
                        question_image = None


                    if answer_size == "0":
                        messages.error(request,"Select Answer Field!")
                        return redirect('users:edit_exam',course_id,section_exam.pk)

                    if box_height == 0:
                        messages.error(request,"Box Height cannot be 0!")
                        return redirect('users:edit_exam',course_id,section_exam.pk)
                    
                    result = Save.save_question_for_exam(exam_id,question_set,question,answer_size,marks,box_height,question_image,question_pk)
                    if result[0]:
                        messages.success(request,result[1])
                        return redirect('users:edit_exam',course_id,section_exam.pk)
                    else:
                        messages.error(request,"Could not update! Try again later")
                        return redirect('users:edit_exam',course_id,section_exam.pk)
                    
                if request.POST.get('delete_section_exam'):
                        
                        flag = str(request.POST.get('delete_section_exam'))
        
                        if flag == "1":
                            if Delete.delete_section_exam(exam_id):
                                messages.success(request,"Deleted Successfully!")
                                return redirect('users:course',course_id)
                            else:
                                messages.error(request,"Could not delete! Try again later")
                                return redirect('users:course',course_id)
                            
                if request.POST.get('save_time'):

                    time_number = request.POST.get('time_number')
                    time_second_number = request.POST.get('time_second_number')

                    time = time_number +" "+  time_second_number
                    section_exam.exam_time = time
                    section_exam.save()
                    messages.success(request,"Time Saved!")
                    return redirect('users:edit_exam',course_id,section_exam.pk)
                
                if request.POST.get('shuffle_papers'):

                    if Save.shuffled_papers(course_id):
                        messages.success(request,"Papers Shuffled!")
                        return redirect('users:edit_exam',course_id,section_exam.pk)
                    else:
                        messages.error(request,"Error Occured")
                        return redirect('users:edit_exam',course_id,section_exam.pk)
                    
                if request.POST.get('save_announcement'):

                    annoucement = request.POST.get('announcement')

                    if Save.save_announcement(logged_in_user,section_exam,annoucement):
                        messages.success(request,"Announcement Posted Successfully!")
                        return redirect('users:edit_exam',course_id,section_exam.pk)
                    else:
                        messages.error(request,"Error Occured")
                        return redirect('users:edit_exam',course_id,section_exam.pk)
                
                if request.POST.get('delete_announcement'):

                    annoucement = request.POST.get('announcement')
                    delete = Announcements.objects.get(section_exam = section_exam,announcement = annoucement)
                    delete.delete()

                    messages.success(request,"Announcement Deleted Successfully!")
                    return redirect('users:edit_exam',course_id,section_exam.pk)
                    

                if request.POST.get('download_question') or request.POST.get('answer_template'):
                    
                    section_exam = Load_Courses.get_saved_section_exams(exam_id)
                    doc = Document()

                    # Set page margins
                    if section_exam.exam_mode.mode_id == 1:
                        set_number = request.POST.get('question_set')
                        questions = Load_Courses.get_questions_and_marks_list(exam_id,set_number)
                        
                        for section in doc.sections:
                            section.top_margin = Inches(1)  # Adjust margin as needed
                            section.bottom_margin = Inches(1)
                            section.left_margin = Inches(1)
                            section.right_margin = Inches(1)

                    
                        doc.add_picture('graduation (1).png', width=Inches(2)) 
                        last_paragraph = doc.paragraphs[-1]
                        last_paragraph.alignment = 1  # Center alignment   
                        doc.add_paragraph() 
                        # Add Exam type, Name, ID, Section
                        x = doc.add_heading(f'{section_exam.exam_type}', level=1)
                        x.alignment = 1
                        for run in x.runs:
                            run.font.color.rgb = None  # Clear any existing color
                            run.font.color.theme_color = 0  # Set the font color to black
                        doc.add_paragraph()
                        
                        x = doc.add_heading('Name:', level=1)

                        for run in x.runs:
                            run.font.color.rgb = None  # Clear any existing color
                            run.font.color.theme_color = 0  # Set the font color to black
                        x = doc.add_heading('ID:', level=1)
                        for run in x.runs:
                            run.font.color.rgb = None  # Clear any existing color
                            run.font.color.theme_color = 0  # Set the font color to black
                        x = doc.add_heading('Section:', level=1)

                        doc.add_paragraph()
                        doc.add_paragraph()
                        doc.add_paragraph()

                        for run in x.runs:
                            run.font.color.rgb = None  
                            run.font.color.theme_color = 0 

                        x =  doc.add_heading('Title:', level=1)
                        for run in x.runs:
                            run.font.color.rgb = None 
                            run.font.color.theme_color = 0  

                        # Add Date
                        x = doc.add_heading('Semester:', level=1)
                        for run in x.runs:
                            run.font.color.rgb = None
                            run.font.color.theme_color = 0  

                        # Add Date
                        x = doc.add_heading('Date:', level=1)
                        for run in x.runs:
                            run.font.color.rgb = None
                            run.font.color.theme_color = 0 

                        # Add Date
                        x = doc.add_heading('Signature:', level=1)
                        for run in x.runs:
                            run.font.color.rgb = None  
                            run.font.color.theme_color = 0  

                        # Add a page break to start the main content on a new page
                        doc.add_page_break()
                        p = doc.add_paragraph()
                        x = p.add_run(f"Set {set_number}")
                        x.bold = True
                        doc.add_paragraph()
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


                        # Iterate through questions and answers
                        for q, m, l, img in zip(questions[0], questions[1],questions[2],questions[3]):
                            # Add the question
                            p = doc.add_paragraph()
                            p.add_run("Q. ").bold=True
                            p.add_run(q).bold = True
                            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                            
                            p = doc.add_paragraph()
                            p.add_run(f"[{m}]")
                            p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            if img:
                                doc.add_picture(os.path.join(settings.MEDIA_URL,img.path),height = Inches(3))
                                last_paragraph = doc.paragraphs[-1]
                                last_paragraph.alignment = 1
                    
                            table = doc.add_table(rows=1, cols=1)
                            table.autofit = False
                            table.columns[0].width = Pt(50)
                            table.rows[0].height = Pt(l) 
                            
                            cell = table.cell(0, 0)
                            tc_pr = cell._element.get_or_add_tcPr()
                            borders = OxmlElement('w:tcBorders')
                            tc_pr.append(borders)
                            
                            for border_type in ['top', 'left', 'bottom', 'right']:
                                border_elm = OxmlElement(f'w:{border_type}')
                                border_elm.set(qn('w:val'), 'dotted')
                                border_elm.set(qn('w:sz'), '15')
                                border_elm.set(qn('w:space'), '0')
                                border_elm.set(qn('w:color'), '000000')
                                border_elm.set(qn('w:rounded'), 'true') 
                                borders.append(border_elm)

                            # # Add nested table in the top right corner
                            # nested_table = cell.add_table(rows=2, cols=2)
                            # nested_table.columns[0].width = Inches(0.5)
                            # nested_table.columns[1].width = Inches(0.5)
                            # nested_table.rows[0].height = Pt(5)
                            # nested_table.rows[1].height = Pt(5)

                            # # Add content to the top right cell of the nested table
                            # top_right_cell = nested_table.cell(0, 1)
                            # p = top_right_cell.add_paragraph()
                            # p.add_run("Small Box").bold = True
                            # p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                            
                            # # Add borders to the nested table cells
                            # for row in nested_table.rows:
                                
                            #     for cell in row.cells:
                            #         tc_pr = cell._element.get_or_add_tcPr()
                            #         borders2 = OxmlElement('w:tcBorders')
                            #         tc_pr.append(borders2)
                            #         for border_type in ['top', 'left', 'bottom', 'right']:
                            #             border = OxmlElement(f'w:{border_type}')
                            #             border.set(qn('w:val'), 'single')
                            #             border.set(qn('w:sz'), '15')
                            #             border.set(qn('w:space'), '0')
                            #             border.set(qn('w:color'), '000000')
                            #             border.set(qn('w:rounded'), 'true') 
                            #             borders2.append(border)
                            
                            doc.add_paragraph()
                            doc.add_paragraph()
                        # Save the document to a BytesIO object
                        doc_stream = BytesIO()
                        doc.save(doc_stream)
                        doc_stream.seek(0)

                        # Return the Word document as an attachment
                        response = HttpResponse(doc_stream, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                        response['Content-Disposition'] = f'attachment; filename=Set:{set_number}_{section_exam.section.course_id.course_name}.{section_exam.section.section_number}.docx'
                        return response
                        
                    elif section_exam.exam_mode.mode_id == 3:

                        set_number = "A"
                        print("herer")
                        questions = Load_Courses.get_questions_and_marks_list(exam_id,set_number)
                        
                        for section in doc.sections:
                            section.top_margin = Inches(1)  # Adjust margin as needed
                            section.bottom_margin = Inches(1)
                            section.left_margin = Inches(1)
                            section.right_margin = Inches(1)



                        # Iterate through questions and answers
                        for q, m, l, img in zip(questions[0], questions[1],questions[2],questions[3]):
                                                
                            table = doc.add_table(rows=1, cols=1)
                            table.autofit = False
                            table.columns[0].width = Pt(50)
                            table.rows[0].height = Pt(l) 
                            
                            cell = table.cell(0, 0)
                            tc_pr = cell._element.get_or_add_tcPr()
                            borders = OxmlElement('w:tcBorders')
                            tc_pr.append(borders)
                            
                            for border_type in ['top', 'left', 'bottom', 'right']:
                                border_elm = OxmlElement(f'w:{border_type}')
                                border_elm.set(qn('w:val'), 'dotted')
                                border_elm.set(qn('w:sz'), '15')
                                border_elm.set(qn('w:space'), '0')
                                border_elm.set(qn('w:color'), '000000')
                                border_elm.set(qn('w:rounded'), 'true') 
                                borders.append(border_elm)
                            
                            doc.add_paragraph()
                            doc.add_paragraph()

                        # Save the document to a BytesIO object
                        doc_stream = BytesIO()
                        doc.save(doc_stream)
                        doc_stream.seek(0)

                        # Return the Word document as an attachment
                        response = HttpResponse(doc_stream, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                        response['Content-Disposition'] = f'attachment; filename={section_exam.section.course_id.course_name}.{section_exam.section.section_number}.docx'
                        return response
                    
                

            none_set_questions = Load_Courses.get_set_questions_for_none(exam_id)
            sets = Load_Courses.get_set_questions_not_none(exam_id)
            
            context = {
                    'page_title':'Check Mate',
                    'user_type':type_of_logged_in_user,
                    'media_url':settings.MEDIA_URL,
                    'logged_in_user':logged_in_user,
                    'year':datetime.now().year,
                    'current_semester':current_semester,
                    
                    'course_id':course_id,
                    'exam_modes':exam_modes,
                    'exam_types':exam_type,
                    'edit_exam':True,
                    'section_exam':section_exam,
                    'question_set':sets,
                    'none_set_questions':none_set_questions,
                    'time':time,
                    'announce':announce,
                }

            return render(request,"add_exam.html",context)
        else:
            context = {
                    'page_title':'Check Mate',
                    'user_type':type_of_logged_in_user,
                    'media_url':settings.MEDIA_URL,
                    'logged_in_user':logged_in_user,
                    'year':datetime.now().year,
                    'current_semester':current_semester,
            }
            return render(request,"access_denied.html",context)
    except Exception as e:
        #saving error information in database if error occured
        logger.error("An error occurred for during logging in at {datetime}".format(datetime=datetime.now()), exc_info=True)
        ErrorHandling.save_system_errors(user,error_name=e,error_traceback=traceback.format_exc())
        return HttpResponse("Bad Request") 
    
@login_required
def exam(request,course_id,exam_type,exam_id,student_id = None):

    try:
        #loading the data to pass them in dictionary, context
        type_of_logged_in_user = Login.user_type_logged_in(request)
        logged_in_user = Login.logged_in_user(request)
        current_semester = Session.objects.get(current=True)
        section_exam = Load_Courses.get_saved_section_exams(exam_id)
        sets = Load_Courses.load_set_of_student(logged_in_user,course_id)
        questions = Load_Courses.get_question_and_marks(exam_id,sets)
        students_exam_material_during_exam = Load_Courses.get_exam_uploaded_students(exam_id)
        answers = Load_Courses.load_students_answer(exam_id,student_id)
        print(questions)
        
        user_submitted = None
        logged_in_ta = False
        ta_allowed = False
        if logged_in_user == None:
            user = request.user.username
            is_uploaded = False
        else:
            user = logged_in_user.user_id
            try:
                user_submitted = Exam_Submitted.objects.get(exam_of = section_exam,student = logged_in_user)
                is_uploaded = user_submitted.is_uploaded
                print(is_uploaded)
            except:
                #ta trying to get in
                ta_allowed = section_exam.ta_available
                is_uploaded = False

        submitted_by = section_exam.section.teaching_assistant.all()
        ta = submitted_by[0].teaching_id

        students = section_exam.section.students.all()
        pending_students = []
        not_pending_students = []
        for student in students:
            stud = student.student_id
            user_submitteds = Exam_Submitted.objects.get(student = stud,exam_of = section_exam)
            if not user_submitteds.is_uploaded:
                pending_students.append(stud)
            else:
                not_pending_students.append(stud)

        if ta == logged_in_user:
            logged_in_ta = True


        if section_exam.is_started or 1 in type_of_logged_in_user or 3 in type_of_logged_in_user or is_uploaded or logged_in_ta or section_exam.exam_mode.mode_id == 1:
            
    
            if request.method == "POST":

                if request.POST.get('shuffle_papers'):

                    if Save.shuffled_papers(course_id):
                        messages.success(request,"Papers Shuffled!")
                        return redirect("users:exam",course_id,exam_type,exam_id)
                    else:
                        messages.error(request,"Error Occured")
                        return redirect("users:exam",course_id,exam_type,exam_id)

                if request.POST.get('start_exam'):

                    if Save.start_exam(exam_id):
                        if Save.shuffled_papers(course_id):
                            messages.success(request,"Exam Started!")
                            return redirect("users:exam",course_id,exam_type,exam_id)
                    else:
                        messages.error(request,"Failed to Start the Exam")
                        return redirect("users:exam",course_id,exam_type,exam_id)

                if request.POST.get('stop_exam'):

                    section_exam = Load_Courses.get_saved_section_exams(exam_id)
                    section_exam.is_stopped=True
                    section_exam.is_started = False
                    section_exam.save()
                    messages.success(request,"Exam Stopped!")
                    return redirect("users:exam",course_id,exam_type,exam_id)

                if request.POST.get('complete_exam'):

                    section_exam = Load_Courses.get_saved_section_exams(exam_id)
                    section_exam.is_completed=True
                    section_exam.save()
                    messages.success(request,"Exam Completed!")
                    return redirect("users:exam",course_id,exam_type,exam_id)

                if request.POST.get('submitCourseFlag'):

                    value = request.POST.get('submitCourseFlag')
                    
                    if value == "1":
                        #logic of ML to separate papers
                        uploaded_file = request.FILES.get('pdf_file')

                    if section_exam.exam_mode.mode_id == 3:
                        if logged_in_ta or 1 in type_of_logged_in_user:
                            messages.error(request,"As an Intructor or TA you can't upload!")
                            return redirect("users:exam",course_id,exam_type,exam_id)
                        else:
                            if Save.uploaded_answer_file(logged_in_user,uploaded_file,exam_id):
                                messages.success(request,"File Uploaded!")
                                #TODO:Redirect to page where uploaded file i viewed
                                return redirect("users:exam",course_id,exam_type,exam_id)
                            else:
                                messages.error(request,"Error Occured while uploaded!")
                                return redirect("users:exam",course_id,exam_type,exam_id)
                    elif section_exam.exam_mode.mode_id == 1:
                        if Save.uploaded_answer_file(logged_in_user,uploaded_file,exam_id,student_id):
                            messages.success(request,"File Uploaded!")
                            #TODO:Redirect to page where uploaded file i viewed
                            return redirect("users:exam",course_id,exam_type,exam_id,student_id)
                        else:
                            messages.error(request,"Error Occured while uploaded!")
                            return redirect("users:exam",course_id,exam_type,exam_id,student_id)



               
            context = {
                        'page_title':'Check Mate',
                        'user_type':type_of_logged_in_user,
                        'media_url':settings.MEDIA_URL,
                        'logged_in_user':logged_in_user,
                        'year':datetime.now().year,
                        'current_semester':current_semester,

                        'questions':questions[0],
                        'total_marks':questions[1],
                        'type_of_logged_in_user':type_of_logged_in_user,
                        'section_exam':section_exam,
                        'exam_type':exam_type,
                        'students_exam_material_during_exam':students_exam_material_during_exam,
                        'course_id':course_id,
                        'exam_id':exam_id,
                        'user_submitted':user_submitted,
                        'ta_allowed':ta_allowed,
                        'logged_in_ta':logged_in_ta,
                        'all_students':pending_students,
                        'not_pending':not_pending_students,
                        'student_id':student_id,
                        'is_uploaded':is_uploaded,
                        'quest_ans_dic':answers,
            }
            return render(request,"exam.html",context)
        else:
            context = {
                    'page_title':'Check Mate',
                    'user_type':type_of_logged_in_user,
                    'media_url':settings.MEDIA_URL,
                    'logged_in_user':logged_in_user,
                    'year':datetime.now().year,
                    'current_semester':current_semester,
            }
            return render(request,"access_denied.html",context)
        
    except Exception as e:
        #saving error information in database if error occured
        logger.error("An error occurred for during logging in at {datetime}".format(datetime=datetime.now()), exc_info=True)
        ErrorHandling.save_system_errors(user,error_name=e,error_traceback=traceback.format_exc())
        return HttpResponse("Bad Request")
@login_required
def review_paper_all(request,course_id,exam_id):

    try:
        #loading the data to pass them in dictionary, context
        type_of_logged_in_user = Login.user_type_logged_in(request)
        logged_in_user = Login.logged_in_user(request)
        current_semester = Session.objects.get(current=True)
        all_students_papers = Load_Courses.get_all_students_submission(exam_id)
        section_exam = Load_Courses.get_saved_section_exams(exam_id)
        
        logged_in_ta = False
        submitted_by = section_exam.section.teaching_assistant.all()
        ta = submitted_by[0].teaching_id
            
        if ta == logged_in_user:
            logged_in_ta = True

        if logged_in_user == None:
            user = request.user.username
        else:
            user = logged_in_user.user_id

        if 1 in type_of_logged_in_user or 3 in type_of_logged_in_user or logged_in_ta:
            context = {
                            'page_title':'Check Mate',
                            'user_type':type_of_logged_in_user,
                            'media_url':settings.MEDIA_URL,
                            'logged_in_user':logged_in_user,
                            'year':datetime.now().year,
                            'current_semester':current_semester,

                            'all_students_papers':all_students_papers,
                            'course_id':course_id,
                            'exam_id':exam_id,
                            'logged_in_ta':logged_in_ta,
            }

            return render(request,"review_paper.html",context)
        else:
            context = {
                    'page_title':'Check Mate',
                    'user_type':type_of_logged_in_user,
                    'media_url':settings.MEDIA_URL,
                    'logged_in_user':logged_in_user,
                    'year':datetime.now().year,
                    'current_semester':current_semester,
            }
            return render(request,"access_denied.html",context)
        
    except Exception as e:
        #saving error information in database if error occured
        logger.error("An error occurred for during logging in at {datetime}".format(datetime=datetime.now()), exc_info=True)
        ErrorHandling.save_system_errors(user,error_name=e,error_traceback=traceback.format_exc())
        return HttpResponse("Bad Request")

@login_required
def student_paper(request,course_id,exam_id,student_id):

    try:
        #loading the data to pass them in dictionary, context
        type_of_logged_in_user = Login.user_type_logged_in(request)
        logged_in_user = Login.logged_in_user(request)
        current_semester = Session.objects.get(current=True)
        question_answer = Load_Courses.get_question_and_answer_of_student(exam_id,student_id)
        section_exam = Load_Courses.get_saved_section_exams(exam_id)
        
        logged_in_ta = False
        submitted_by = section_exam.section.teaching_assistant.all()
        ta = submitted_by[0].teaching_id
            
        if ta == logged_in_user:
            logged_in_ta = True
        
        
        if logged_in_user == None:
            user = request.user.username
        else:
            user = logged_in_user.user_id

        allowed_user = None

        if question_answer[1] == logged_in_user:
            allowed_user = True
        else:
            allowed_user = False
        
        if 1 in type_of_logged_in_user or 3 in type_of_logged_in_user or allowed_user or logged_in_ta:
        
            context = {
                                'page_title':'Check Mate',
                                'user_type':type_of_logged_in_user,
                                'media_url':settings.MEDIA_URL,
                                'logged_in_user':logged_in_user,
                                'year':datetime.now().year,
                                'current_semester':current_semester,
                                'course_id':course_id,
                                'exam_id':exam_id,
                                'student_id':student_id,

                                'question_answers':question_answer[0],
                                'user':question_answer[1],
                                'total_marks':question_answer[2],
                                'obtained_marks':question_answer[3],
                                'set_number':question_answer[4],
                                'first_question':question_answer[5],
                                'logged_in_ta':logged_in_ta,
                }
        
            return render(request,"student_exam_submit_page.html",context)
        else:
            context = {
                    'page_title':'Check Mate',
                    'user_type':type_of_logged_in_user,
                    'media_url':settings.MEDIA_URL,
                    'logged_in_user':logged_in_user,
                    'year':datetime.now().year,
                    'current_semester':current_semester,
            }
            return render(request,"access_denied.html",context)
    except Exception as e:
        #saving error information in database if error occured
        logger.error("An error occurred for during logging in at {datetime}".format(datetime=datetime.now()), exc_info=True)
        ErrorHandling.save_system_errors(user,error_name=e,error_traceback=traceback.format_exc())
        return HttpResponse("Bad Request")
    
@login_required
def paper_view(request,course_id,exam_id,student_id,question_pk):

    try:
        #loading the data to pass them in dictionary, context
        type_of_logged_in_user = Login.user_type_logged_in(request)
        logged_in_user = Login.logged_in_user(request)
        current_semester = Session.objects.get(current=True)
        section_exam = Load_Courses.get_saved_section_exams(exam_id)
        logged_in_ta = False

        question = Question.objects.get(pk = question_pk,questions_of = section_exam)
        print(question.pk)
        answer = Answer.objects.get(answer_of = question,uploaded_by = School_Users.objects.get(user_id = student_id))


        try:
            question_next = Question.objects.filter(pk__gt=question.pk,questions_of = section_exam,question_set = question.question_set).order_by('pk').first()
            answer_next = Answer.objects.get(answer_of = question,uploaded_by = School_Users.objects.get(user_id = student_id))
        except:
            question_next = None
            answer_next = None

        try:
            question_back= Question.objects.filter(pk__lt=question.pk,questions_of = section_exam,question_set = question.question_set).order_by('-pk').first()
            answer_back = Answer.objects.get(answer_of = question,uploaded_by = School_Users.objects.get(user_id = student_id))
        except:
            question_back = None
            answer_back = None

        if logged_in_user == None:
            user = request.user.username
        else:
            user = logged_in_user.user_id

        submitted_by = section_exam.section.teaching_assistant.all()
        ta = submitted_by[0].teaching_id
            
        if ta == logged_in_user:
            logged_in_ta = True

        if request.method == 'POST':

            annotated_image = request.POST.get('annotated_image')
            # Process the annotated image data here
            # Example: Save the image to a file or database


            if annotated_image:
                path = settings.MEDIA_ROOT+str(answer.answer_image)
                if os.path.isfile(path):
                    print("Here")
                    os.remove(path)

                format, imgstr = annotated_image.split(';base64,')
        
                # Decode the base64 string to binary data
                img_data = base64.b64decode(imgstr)
                
                # Generate a file name
                file_name = f'annotated_{int(time.time())}.png'
                
                # Create a Django file
                django_file = ContentFile(img_data, name=file_name)
                
                # Assign this file to the answer_image field
                answer.answer_image.save(file_name, django_file, save=False)
                answer.save()
            

            if request.POST.get('submit_marks_comment'):
                marks = request.POST.get('marks')
                comment = request.POST.get('comment')

                if Save.save_marks_comment(question_pk,marks,comment,student_id):
                    messages.success(request,"Saved!")
                    #TODO:Redirect to page where uploaded file i viewed
                    return redirect("users:paper_view",course_id,exam_id,student_id,question_pk)
                else:
                    messages.error(request,"Error Occured while saving!")
                    return redirect("users:exam",course_id,exam_id,student_id,question_pk)




        context = {
                                'page_title':'Check Mate',
                                'user_type':type_of_logged_in_user,
                                'media_url':settings.MEDIA_URL,
                                'logged_in_user':logged_in_user,
                                'year':datetime.now().year,
                                'current_semester':current_semester,
                                'logged_in_ta':logged_in_ta,

                                'question':question,
                                'answer':answer,
                                'question_next':question_next,
                                'answer_next':answer_next,
                                'section_exam':section_exam,
                                'course_id':course_id,
                                'exam_id':exam_id,
                                'student_id':student_id,
                                'question_back':question_back,
                                'answer_back':answer_back,
                                'question_pk':question_pk,

        }

        return render(request,"paper_view.html",context)
    except Exception as e:
        #saving error information in database if error occured
        logger.error("An error occurred for during logging in at {datetime}".format(datetime=datetime.now()), exc_info=True)
        ErrorHandling.save_system_errors(user,error_name=e,error_traceback=traceback.format_exc())
        return HttpResponse("Bad Request")

@login_required
def generate_spreadsheet(request,course_id):
    try:

        #loading the data to pass them in dictionary, context
        type_of_logged_in_user = Login.user_type_logged_in(request)
        logged_in_user = Login.logged_in_user(request)
        current_semester = Session.objects.get(current=True)
        course_section = Load_Courses.get_specific_course_section(course_id)

        midterm = 0
        
        numbers = Load_Courses.number_of_quizzes_and_midterm(course_id)
        try:
            midterm = 35/len(numbers[1])
        except:
            midterm = 0

       
        if logged_in_user == None:
            user = request.user.username
        else:
            user = logged_in_user.user_id

        

        if request.method == "POST":

            if request.POST.get('generate_excel'):
                
                midterm_weights = []
                mid_number = 0
                best_score_number = request.POST.get('best_score')
                # quiz_weight = int(request.POST.get('quiz_weight'))
                # if quiz_weight == None:
                #     quiz_weight = 15
                quizzes_number = best_score_number

                if best_score_number == "All":
                    best_score_number = len(numbers[0])
                else:
                    best_score_number = int(best_score_number)

                # for i in range(len(numbers[1])):
                #     weight = int(request.POST.get(f"midterm_weight_{i+1}"))
                #     midterm_weights.append(weight)
                #     mid_number += 1 
                
                # try:
                #     final_weight = int(request.POST.get('final_weight'))
                # except:
                #     final_weight = 0

                result = Load_Courses.load_spreadsheet_info(course_id,quizzes_number)

                semester = Session.objects.get(current=True)
                date=datetime.now()
                response = HttpResponse(
                    content_type='application/ms-excel')  # eclaring content type for the excel files
                response['Content-Disposition'] = f'attachment; filename={course_section.course_id.course_name}.{course_section.section_number} - {semester}: Semester Grade Sheet.xls'  # making files downloadable with name of session and timestamp
                # adding encoding to the workbook
                workBook = xlwt.Workbook(encoding='utf-8')
                # opening an worksheet to work with the columns
                workSheet = workBook.add_sheet(f'Semester Grade Sheet of {semester}')

                # generating the first row
                row_num = 0
                font_style = xlwt.XFStyle()
                font_style.font.bold = True

                quiz_columns = [f'Quiz-{i}' for i in range(1, best_score_number + 1)]
                mid_columns = [f'MidTerm-{i}' for i in range(1,len(numbers[1])+1)]

                # Defining columns that will stay in the first row
                columns = ['ID', 'Name'] + quiz_columns + ['Quiz (Average)'] + mid_columns + ['Mid (Average)', 'Final']

                # Defining first column
                column_widths = [5000,5000]
                column_widths+=[5000]*best_score_number
                column_widths+=[5000]
                column_widths+=[5000]*mid_number
                column_widths+=[5000,5000]

                for col, width in enumerate(column_widths):
                    workSheet.col(col).width = width


                for column in range(len(columns)):
                    workSheet.write(row_num, column, columns[column], xlwt.easyxf('align: horiz center; alignment: wrap True'))

                # reverting font style to default
                font_style = xlwt.XFStyle()

                # Center alignment style
                center_alignment = xlwt.easyxf('align: horiz center')
                # Word wrap style
                word_wrap_style = xlwt.easyxf('alignment: wrap True')
                row_num += 1
                # events= Branch.load_all_inter_branch_collaborations_with_events_yearly(year,1)
                # sl_num = 0
                for student, scores in result.items():
                    # Write student ID and name
                    workSheet.write(row_num, 0, student.user_id)
                    workSheet.write(row_num, 1, student.user_first_name)

                    # Write quiz scores
                    for idx, quiz_score in enumerate(scores['Quizzes']):
                        workSheet.write(row_num, 2 + idx, quiz_score)

                    # Write quiz average
                    workSheet.write(row_num, 2 + len(quiz_columns), scores['Quiz Average'])

                    # Write midterm scores
                    for idx, mid_score in enumerate(scores['Mids']):
                        workSheet.write(row_num, 3 + len(quiz_columns) + idx, mid_score)

                    # Write midterm average
                    workSheet.write(row_num, 3 + len(quiz_columns) + len(mid_columns), scores['Midterm Average'])

                    # Write final score
                    workSheet.write(row_num, 4 + len(quiz_columns) + len(mid_columns), scores['Final Average'])
                    workSheet.write(row_num, 5 + len(quiz_columns) + len(mid_columns), scores['result'])
                    # Move to the next row
                    row_num += 1
                # for event,collaborations in events.items():
                #     row_num += 1
                #     sl_num += 1
                #     workSheet.write(row_num,0 , sl_num,  center_alignment)
                #     workSheet.write(row_num,1 , event.event_name,  center_alignment)
                #     workSheet.write(row_num,2 , event.event_date.strftime('%Y-%m-%d'),  center_alignment)
                #     workSheet.write(row_num,3 , event.event_organiser.group_name,  center_alignment)
                #     collaborations_text = ""
                #     for collabs in collaborations:
                #         collaborations_text += collabs + '\n'
                #     workSheet.write(row_num, 4, collaborations_text, word_wrap_style) 
                #     categories = ""   
                #     for event_type in event.event_type.all():
                #         categories+=event_type.event_category + '\n'  
                #     workSheet.write(row_num, 5, categories, word_wrap_style)
                #     venue_list = Branch.get_selected_venues(event.pk)
                #     venues=""
                #     for venue in venue_list:
                #         venues += venue + '\n'
                #     workSheet.write(row_num, 6, venues, word_wrap_style)
                        
                workBook.save(response)
                return (response)

                



        context = {
                    'page_title':'Check Mate',
                    'user_type':type_of_logged_in_user,
                    'media_url':settings.MEDIA_URL,
                    'logged_in_user':logged_in_user,
                    'year':datetime.now().year,
                    'current_semester':current_semester,

                    'course_id':course_id,
                    'quiz_number':numbers[0],
                    'midterm_number':numbers[1],
                    'final':numbers[2],
                    'midterm':midterm,
        }

        return render(request,"excel_creation_page.html",context)

    except Exception as e:
        #saving error information in database if error occured
        logger.error("An error occurred for during logging in at {datetime}".format(datetime=datetime.now()), exc_info=True)
        ErrorHandling.save_system_errors(user,error_name=e,error_traceback=traceback.format_exc())
        return HttpResponse("Bad Request")

@login_required
def announcements(request):

    try:
        #loading the data to pass them in dictionary, context
        type_of_logged_in_user = Login.user_type_logged_in(request)
        logged_in_user = Login.logged_in_user(request)
        current_semester = Session.objects.get(current=True)

        if logged_in_user == None:
            user = request.user.username
        else:
            user = logged_in_user.user_id

        announcements = Load_Courses.load_announcements(logged_in_user,type_of_logged_in_user)
        print(announcements)

        if request.method == "POST":

            if request.POST.get('download_paper'):
                
                doc = Document()
                section_exam = int(request.POST.get('get_value'))
                set_number = "A"
                questions = Load_Courses.get_questions_and_marks_list(section_exam,set_number)
                for section in doc.sections:
                    section.top_margin = Inches(1)  # Adjust margin as needed
                    section.bottom_margin = Inches(1)
                    section.left_margin = Inches(1)
                    section.right_margin = Inches(1)



                # Iterate through questions and answers
                for q, m, l, img in zip(questions[0], questions[1],questions[2],questions[3]):
                                        
                    table = doc.add_table(rows=1, cols=1)
                    table.autofit = False
                    table.columns[0].width = Pt(50)
                    table.rows[0].height = Pt(l) 
                    
                    cell = table.cell(0, 0)
                    tc_pr = cell._element.get_or_add_tcPr()
                    borders = OxmlElement('w:tcBorders')
                    tc_pr.append(borders)
                    
                    for border_type in ['top', 'left', 'bottom', 'right']:
                        border_elm = OxmlElement(f'w:{border_type}')
                        border_elm.set(qn('w:val'), 'dotted')
                        border_elm.set(qn('w:sz'), '15')
                        border_elm.set(qn('w:space'), '0')
                        border_elm.set(qn('w:color'), '000000')
                        border_elm.set(qn('w:rounded'), 'true') 
                        borders.append(border_elm)
                    
                    doc.add_paragraph()
                    doc.add_paragraph()

                # Save the document to a BytesIO object
                doc_stream = BytesIO()
                doc.save(doc_stream)
                doc_stream.seek(0)

                # Return the Word document as an attachment
                response = HttpResponse(doc_stream, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                response['Content-Disposition'] = f'attachment; filename=Answer Template.docx'
                return response


        context = {
                    'page_title':'Check Mate',
                    'user_type':type_of_logged_in_user,
                    'media_url':settings.MEDIA_URL,
                    'logged_in_user':logged_in_user,
                    'year':datetime.now().year,
                    'current_semester':current_semester,
                    'announcements':announcements,
        }
        return render(request,"announcements.html",context)

    except Exception as e:
        #saving error information in database if error occured
        logger.error("An error occurred for during logging in at {datetime}".format(datetime=datetime.now()), exc_info=True)
        ErrorHandling.save_system_errors(user,error_name=e,error_traceback=traceback.format_exc())
        return HttpResponse("Bad Request")
